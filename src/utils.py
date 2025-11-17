"""
Funções auxiliares para cálculos e processamento
"""
import pandas as pd
from datetime import datetime, date
import os


def calcular_dias_manutencao(data_entrada, data_saida):
    """
    Calcula dias em manutenção
    Se data_saida vazia, usa hoje
    """
    if pd.isna(data_entrada) or data_entrada == '' or data_entrada is None:
        return 0
    
    try:
        data_entrada = pd.to_datetime(data_entrada)
    except:
        return 0
    
    if pd.isna(data_saida) or data_saida == '' or data_saida is None:
        data_fim = datetime.now()
    else:
        try:
            data_fim = pd.to_datetime(data_saida)
        except:
            data_fim = datetime.now()
    
    dias = (data_fim - data_entrada).days
    return max(0, dias)


def calcular_status(data_entrada, data_saida, status_atual=''):
    """
    Calcula status dinâmico baseado nas datas
    """
    if pd.notna(data_entrada) and data_entrada != '' and (pd.isna(data_saida) or data_saida == ''):
        return 'EM SERVIÇO'
    elif pd.notna(data_entrada) and data_entrada != '' and pd.notna(data_saida) and data_saida != '':
        return 'FINALIZADO'
    else:
        return status_atual if status_atual else ''


def validar_data(data_str):
    """
    Valida e converte string para data
    """
    if not data_str or data_str == '':
        return None
    
    try:
        if isinstance(data_str, str):
            # Tenta vários formatos
            for fmt in ['%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y']:
                try:
                    return pd.to_datetime(data_str, format=fmt)
                except:
                    continue
        return pd.to_datetime(data_str)
    except:
        return None


def formatar_data_br(data):
    """
    Formata data para padrão brasileiro
    """
    if pd.isna(data) or data is None or data == '':
        return ''
    
    if isinstance(data, str):
        return data
    
    try:
        return data.strftime('%d/%m/%Y')
    except:
        return str(data)


def validar_numero(valor, padrao=0):
    """
    Valida e converte para número
    """
    if valor is None or valor == '':
        return padrao
    
    try:
        # Remove pontos e vírgulas
        valor_str = str(valor).replace('.', '').replace(',', '')
        return int(valor_str)
    except:
        return padrao


def limpar_texto(texto):
    """
    Remove espaços extras e normaliza texto
    """
    if pd.isna(texto) or texto is None:
        return ''
    return str(texto).strip()


def gerar_relatorio_pdf(dados, estatisticas, arquivo_saida):
    """
    Gera relatório em PDF
    """
    try:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        # Criar documento
        doc = SimpleDocTemplate(arquivo_saida, pagesize=landscape(A4))
        elements = []
        styles = getSampleStyleSheet()
        
        # Título
        titulo_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        elements.append(Paragraph('RELATÓRIO DE MANUTENÇÃO - ALS', titulo_style))
        elements.append(Spacer(1, 0.5*cm))
        
        # Data do relatório
        data_relatorio = datetime.now().strftime('%d/%m/%Y %H:%M')
        elements.append(Paragraph(f'Gerado em: {data_relatorio}', styles['Normal']))
        elements.append(Spacer(1, 0.5*cm))
        
        # Estatísticas
        stats_data = [
            ['ESTATÍSTICAS GERAIS'],
            ['Total de Registros:', str(estatisticas.get('total', 0))],
            ['Em Serviço:', str(estatisticas.get('em_servico', 0))],
            ['Finalizados:', str(estatisticas.get('finalizados', 0))],
            ['Tempo Médio (dias):', f"{estatisticas.get('tempo_medio', 0):.1f}"],
            ['Placas Únicas:', str(estatisticas.get('placas_unicas', 0))]
        ]
        
        stats_table = Table(stats_data, colWidths=[8*cm, 4*cm])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(stats_table)
        elements.append(Spacer(1, 1*cm))
        
        # Tabela de dados
        if not dados.empty:
            elements.append(Paragraph('REGISTROS DE MANUTENÇÃO', titulo_style))
            elements.append(Spacer(1, 0.5*cm))
            
            # Preparar dados da tabela
            colunas = ['PLACA', 'VEÍCULO', 'DATA ENTRADA', 'DATA SAÍDA', 'DIAS', 'STATUS']
            table_data = [colunas]
            
            for _, row in dados.head(50).iterrows():  # Limite 50 registros por página
                table_data.append([
                    str(row.get('PLACA', '')),
                    str(row.get('VEÍCULO', ''))[:15],
                    formatar_data_br(row.get('DATA ENTRADA', '')),
                    formatar_data_br(row.get('DATA SAÍDA', '')),
                    str(row.get('TOTAL DE DIAS EM MANUTENÇÃO', 0)),
                    str(row.get('STATUS', ''))
                ])
            
            data_table = Table(table_data, colWidths=[3*cm, 4*cm, 3*cm, 3*cm, 2*cm, 3*cm])
            data_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
            ]))
            elements.append(data_table)
        
        # Gerar PDF
        doc.build(elements)
        return True, arquivo_saida
        
    except Exception as e:
        return False, f"Erro ao gerar PDF: {str(e)}"


def gerar_relatorio_word(dados, estatisticas, arquivo_saida):
    """
    Gera relatório em Word (.docx)
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Criar documento
        doc = Document()
        
        # Título
        titulo = doc.add_heading('RELATÓRIO DE MANUTENÇÃO - ALS', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data
        data_relatorio = datetime.now().strftime('%d/%m/%Y %H:%M')
        p = doc.add_paragraph(f'Gerado em: {data_relatorio}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Estatísticas
        doc.add_heading('ESTATÍSTICAS GERAIS', 2)
        table_stats = doc.add_table(rows=6, cols=2)
        table_stats.style = 'Light Grid Accent 1'
        
        stats_data = [
            ['Total de Registros:', str(estatisticas.get('total', 0))],
            ['Em Serviço:', str(estatisticas.get('em_servico', 0))],
            ['Finalizados:', str(estatisticas.get('finalizados', 0))],
            ['Tempo Médio (dias):', f"{estatisticas.get('tempo_medio', 0):.1f}"],
            ['Placas Únicas:', str(estatisticas.get('placas_unicas', 0))],
            ['', '']
        ]
        
        for i, (label, valor) in enumerate(stats_data[:-1]):
            table_stats.rows[i].cells[0].text = label
            table_stats.rows[i].cells[1].text = valor
        
        doc.add_paragraph()
        
        # Tabela de dados
        if not dados.empty:
            doc.add_heading('REGISTROS DE MANUTENÇÃO', 2)
            
            # Criar tabela
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            
            # Cabeçalhos
            hdr_cells = table.rows[0].cells
            headers = ['PLACA', 'VEÍCULO', 'DATA ENTRADA', 'DATA SAÍDA', 'DIAS', 'STATUS']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Dados (limite de 100 registros)
            for _, row in dados.head(100).iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('PLACA', ''))
                row_cells[1].text = str(row.get('VEÍCULO', ''))[:20]
                row_cells[2].text = formatar_data_br(row.get('DATA ENTRADA', ''))
                row_cells[3].text = formatar_data_br(row.get('DATA SAÍDA', ''))
                row_cells[4].text = str(row.get('TOTAL DE DIAS EM MANUTENÇÃO', 0))
                row_cells[5].text = str(row.get('STATUS', ''))
        
        # Salvar
        doc.save(arquivo_saida)
        return True, arquivo_saida
        
    except Exception as e:
        return False, f"Erro ao gerar Word: {str(e)}"
